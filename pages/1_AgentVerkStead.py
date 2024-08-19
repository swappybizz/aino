import streamlit as st
from pymongo import MongoClient
import fitz  # PyMuPDF
import docx
from datetime import datetime
import base64
from agent_helpers import add_agent 

MONGO_URI = st.secrets["mongo_uri"]
client = MongoClient(MONGO_URI)
db = client["aino_db"]
agent_collection = db["agents"]
session_collection = db["client_sessions"]

LOGO_URL_SMALL = "aino.png"
st.logo(image="aino.png",icon_image=LOGO_URL_SMALL) 
st.set_page_config(
    page_title="Assistant Forge Tools",
    page_icon="🔥",
    layout="wide",

)

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
    
if 'loaded_agent' not in st.session_state:
    st.session_state.loaded_agent = None

if "all_Expert_client_ai_sessions" not in st.session_state:
    st.session_state.all_Expert_client_ai_sessions = []
    
if "laoded_expert_client_ai_chat" not in st.session_state:
    st.session_state.laoded_expert_client_ai_chat = []
  
def get_expert_enquiries(username):
    #return the whole sessions collection 
    # find all sessions with parent as username
    print(f" inding all sessions with parent as {username}")
    all_sessions = session_collection.find({"parent": username})
    all_sessions = list(all_sessions)
    print(f"all sessions: {all_sessions}")
    if "all_Expert_client_ai_sessions" not in st.session_state:
        st.session_state.all_Expert_client_ai_sessions = []
    st.session_state.all_Expert_client_ai_sessions = all_sessions
    
    
if st.session_state.logged_in:
    tab1, tab2, tab3 = st.tabs(["AgentClientForge", "ClientDMs","Socials"])
    with tab1:
        maincol1, maincol2 = st.columns([1, 2])
        with maincol1:
            with st.expander("Add New Agent", expanded=True):
                with st.form("Add New Agent",
                             clear_on_submit=False, #TODO: change in production
                             border=False):
                    agent_name = st.text_input("Agent Name")
                    agent_userTitle = st.text_input(
                        "AIM of the Agent",
                        help="e.g. Responsible for educating employees during onboarding.",
                    )
                        # multi-select for type of inputs
                    inputs = st.multiselect(
                            "Type of Inputs",
                            placeholder= "Select what you want from the user",
                            options=["Documents", "QuerySession", "File/s"]
                        )
                    agent_instructions = st.text_area("Agent Instructions")
                        # multi-select for type of deliverables
                    deliverables = st.multiselect(
                            label="Type of Deliverables",
                        options= ["Documents", "UQuerySession","File/s" ],
                        placeholder= "Select what you want to deliver to the user"
                        )
                    uploaded_docs = st.file_uploader(
                        ":red[*Optional] Upload Documents for knowledge &/or Examples etc", type=["pdf", "docx"], accept_multiple_files=True
                    )
                    opening_statement_client_expectation = st.text_area(
                        help="This intiates the conversation, depending on type of Bot you want you can make it enquiring, investigative, informative etc.",
                        label=":red[Cant be modified] : Opening Statement & Client Expectation: Clarify the role of the bot",
                        placeholder="eg. Hello I am XYZ bot I will help you setup a project for Internal Audit for ISO 9001:2015. I will need some information from you to get started. ..."
                    )
                    docs = []
                    
                    # for each file, add filename, filetype, and content
                    
                    for doc in uploaded_docs:
                        if doc.type == "application/pdf":
                            pdf = fitz.open(stream=doc.getvalue(), filetype="pdf")
                            text = ""
                            for page_num in range(pdf.page_count):
                                text += pdf[page_num].get_text()
                            docs.append(
                                {
                                    "filename": doc.name,
                                    "filetype": "pdf",
                                    "content": text,
                                }
                            )
                        elif doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            docx_file = docx.Document(doc)
                            text = ""
                            for para in docx_file.paragraphs:
                                text += para.text
                            docs.append(
                                {
                                    "filename": doc.name,
                                    "filetype": "docx",
                                    "content": text,
                                }
                            )
                            
                    
                    allow_personal  = st.checkbox("Allow Personal Interactions")
                    if st.form_submit_button("Add Agent", use_container_width=True):
                        # check if the form feilds are filled, if empty warn and dont proceed
                        if not agent_name or not agent_userTitle or not agent_instructions or not deliverables or not opening_statement_client_expectation:
                            st.error("All fields are required.")
                        else:
                        # print the form data

                            add_agent(
                               st.session_state.username,
                               agent_name,
                               inputs,
                                agent_userTitle,
                                agent_instructions,
                               deliverables,
                               docs,
                               opening_statement_client_expectation,
                               allow_personal,
                               datetime.now()
                            )
                            st.success("Agent added successfully!")
                        
                    
            with st.container(border=True,height=675):
                
                # get all agents with the username
                agents = agent_collection.find({"parent": st.session_state.username})
                for i, agent in enumerate(agents):
                    with st.expander(f"Agent {i+1}"):
                        agentcol1, agentcol2, agentcol3, agnetcol4 = st.columns(
                            [2, 1, 1, 1]
                        )
                        with agentcol1:
                            st.write(f"**{agent['agent_name']}**")
                        with agentcol3:
                            if st.button(
                                "Load", key=f"pick_{i}", use_container_width=True
                            ):
                                st.session_state.loaded_agent = agent
                                st.success("Agent loaded successfully!")
                                pass
                        # with agentcol2:
                        #     if st.button(
                        #         "🟩", key=f"status_{i}", use_container_width=True
                        #     ):
                        #         pass
                        with agnetcol4:
                            if st.button(
                                "❌", key=f"delete___{i}", use_container_width=True
                            ):
                                agent_collection.delete_one({"_id": agent["_id"]})
                                # delete the agent from the session state
                                
                                
                                st.success("Agent deleted successfully!")
                                st.rerun()
                                
                            
        with maincol2:
            setcol1, setcol2 = st.columns([2, 1])
            with setcol1:
                # if st.session_state.selected_agent:
                    if st.session_state.loaded_agent:
                        with st.form("Instructions For Agents"):
                        
                        
                            ins = st.session_state.loaded_agent["agent_instructions"]
                        
                            new_instructions = st.text_area(label="Instructions", height=640, value=ins)
                            if st.form_submit_button("Save Instructions", use_container_width=True):
                                # change the instructions in the database, and update the session state
                                agent_collection.update_one(
                                    {"_id": st.session_state.loaded_agent["_id"]},
                                    {"$set": {"agent_instructions": new_instructions}},
                                )
                                st.session_state.loaded_agent["agent_instructions"] = new_instructions
                                st.success("Instructions saved successfully!")
                        
                        
            with setcol2:
                fileupload = st.file_uploader("Upload Document", type=["pdf", "docx"])
                if fileupload:
                    if fileupload.type == "application/pdf":
                        pdf = fitz.open(stream=fileupload.getvalue(), filetype="pdf")
                        text = ""
                        for page_num in range(pdf.page_count):
                            text += pdf[page_num].get_text()
                        doc = {
                            "filename": fileupload.name,
                            "filetype": "pdf",
                            "content": text,
                        }
                    elif fileupload.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        docx_file = docx.Document(fileupload)
                        text = ""
                        for para in docx_file.paragraphs:
                            text += para.text
                        doc = {
                            "filename": fileupload.name,
                            "filetype": "docx",
                            "content": text,
                        }
                    agent_collection.update_one(
                            {"_id": st.session_state.loaded_agent["_id"]},
                            {"$push": {"docs_store": doc}},
                        )
                    # add to the session state
                    st.session_state.loaded_agent["docs_store"].append(doc)
                    
                    st.success("File uploaded successfully!Clear the file from above before another files operations")
                
                with st.container():
                    st.write("Uploaded Files:")
                    if st.session_state.loaded_agent:
                        for file in st.session_state.loaded_agent["docs_store"]:

                            with st.container():
                                with st.expander(f"File {file['filename']} | {file["filetype"]}"):
                                    with st.container(border=True):
                                        st.write("Content:")
                                        st.write(file["content"][:100])
                                if st.button(
                                            "❌", key=f"delete___{file['filename']}", use_container_width=True,
                                        ):
                                    
                                    # find the filein the dB. no need to delete it just print its existence
                                    if file in agent_collection.find_one(
                                        {"_id": st.session_state.loaded_agent["_id"]}
                                    )["docs_store"]:
                                    #   print ( file exists)
                                    #  delete the file from the db
                                        agent_collection.update_one(
                                            {"_id": st.session_state.loaded_agent["_id"]},
                                            {"$pull": {"docs_store": file}},
                                        )
                                        
                                        # delete from the session state
                                        st.session_state.loaded_agent["docs_store"].remove(file)
                                        
                                        st.success("File deleted successfully! your Preview ill get updated soon as well")
                                        st.rerun()  
                                        
                                    
                        
    # with tab2:
    #     "### Client Queries"
    #     # user the username to get get
    #     for i in range (7):
    #         with st.form(border=True, key = f"query_{i}"):
    #             st.write(f"Query {i}")
    #             answer =  "query answer if it exists"
    #             reply = st.text_area("Reply", value = answer)
    #             replcol1, replcol2, replcol3, replcol4, replcol5 = st.columns([1,1,1,1, 3])
    #             with replcol1:
    #                 checkbox = st.checkbox("Add to Global Agent knowledge base")
    #             with replcol2:
    #                 checkbox2 = st.checkbox("Make Public")
    #             with replcol3:
    #                 checkbox3 = st.checkbox("Add chat reply in Agent | Client Chat")
    #             with replcol4:
    #                 open_chat = st.form_submit_button(":blue[Open Manuscript]", use_container_width=True)
    #             with replcol5:
    #                 if st.form_submit_button("Reply",use_container_width=True):
    #                     pass
    #             if st.form_submit_button(":red[Dismiss Query]"):
    #                 pass    
                
    with tab2:
        dmcol1, dmcol2 = st.columns([1, 2]) 
        with dmcol1:
            with st.container(border=True, height=240):
                "Clients"
                for i in range(5):
                    with st.expander(f"DM {i}"):
                        "Some details of the DM"
                        load_project_button = st.button("Load Project", key=f"load_project_{i}",use_container_width=True)           
        
            with st.container(border=True, height=300):
                f"Expert Enquiries for {st.session_state.username}"
                # 
                # get all the client_DM for  sessions where username is the parent
                get_expert_enquiries(st.session_state.username)
                for i in st.session_state.all_Expert_client_ai_sessions:
                    load_project_button = st.button(f"Clinet DM {i['client']}",
                                                    key=f"load_project_{i}",
                                                    use_container_width=True)
                    if load_project_button:
                        st.session_state.laoded_expert_client_ai_chat = i
                        st.success("Client session loaded successfully!")
                        

        
        with dmcol2:
            dmmidcol1, dmmidcol2 = st.columns([1, 2])
            with dmmidcol1:

                with st.container(height=555, border=True):
                    st.write("Project AIM")
                    f"""
                    - Agent Name: {st.session_state.laoded_expert_client_ai_chat['agent_name']}
                    - AIM: {st.session_state.laoded_expert_client_ai_chat['agent_userTitle']}
                    - Deliverables: {st.session_state.laoded_expert_client_ai_chat['deliverables']}
                    - Client : {st.session_state.laoded_expert_client_ai_chat['client']}
                    - Client Goal : {st.session_state.laoded_expert_client_ai_chat['client_goal']}
                    """
                    

            with dmmidcol2:
                with st.container(height=500, border=True):
                    for message in st.session_state.laoded_expert_client_ai_chat["client_DM"]:
                        if 'ai' in message:
                            with st.chat_message("assistant"):
                                st.write(message['ai'])
                        elif 'client' in message:
                            with st.chat_message("user"):
                                st.write(message['client'])
                        elif 'expert' in message:
                            with st.chat_message("expert"):
                                st.write(message['expert'])

                
                # get the ganet id from the loaded agent
                # st.write(st.session_state.laoded_expert_client_ai_chat['agent_id'])
                # in the agents collection, find the agent and append to array

                        
                            
                if expert_reply := st.chat_input("Type your message here") :
                    st.session_state.laoded_expert_client_ai_chat["client_DM"].append({"expert": expert_reply})
                    # update the session in the db
                    # TODO:update all_Expert_client_ai_sessions using the session id
                    # if personal is not allowed, add to the client_AI_chat as well
                    session_collection.update_one(
                        {"_id": st.session_state.laoded_expert_client_ai_chat["_id"]},
                        {"$push": {"client_DM": {"expert": expert_reply}}},
                    )
                    if st.session_state.laoded_expert_client_ai_chat["allow_personal"] != "yes":
                        session_collection.update_one(
                            {"_id": st.session_state.laoded_expert_client_ai_chat["_id"]},
                            {"$push": {"client_AI_chat": {"expert": expert_reply}}},
                        )
                    st.rerun()
                add_to_agent_memory = st.checkbox("Add to Agent Memory")
                # TODO: use agent_id to find and append to agents 'Experience_knowledge_base'
                if add_to_agent_memory:
                    #  get the last two items from tjhe chat and add to the agent memory
                    knowledge = st.session_state.laoded_expert_client_ai_chat["client_DM"][-2:]
                    agent_collection.update_one(
                        {"_id": st.session_state.laoded_expert_client_ai_chat['agent_id']},
                        {"$push": {"Experience_knowledge_base": knowledge}},
                    )
                    st.success("Added to Agent Memory")
                        
                    
                
        with st.expander("Project Deliverables", expanded=True):
            delicol1, delicol2 = st.columns([1, 1])
            
            with delicol1:
                ":gray[Expectations of the user]"
                "The user expects the following deliverables:"
                """
                - :blue[Important Information like Date 1]
                - :blue[Important Information like number of files and quality expected]
                
                
                """
                with st.container(border=True):
                    ":gray[user Feedback and Review]"
                    rating =  st.slider("Rating", 0, 5, 3,disabled=True)
                    "feedback"
                    st.text_area("Feedback",height=100,disabled=True)
                    
                chalcol1, chalcol2 = st.columns([1, 1])
                with chalcol1:
                    st.button("Accept", key="accept", use_container_width=True)
                with chalcol2:
                    st.button(":red[Challenge]", key="reject", use_container_width=True)
                    
                
            with delicol2:
                upload_file = st.file_uploader("Upload File", type=["pdf", "docx"])
                
                with st.container(height=300):
                    "uploaded files"
                    for i in range(3):
                        with st.container(border=True):
                            col1, col2 ,col3 = st.columns([2, 1,1])
                            with col1:
                                st.write(f"File {i}")
                                
                            with col2:
                                view = st.button("↗", key=f"view_file_{i}", use_container_width=True,help="View the file")
                            with col3:
                                
                                if st.button("❌", key=f"delete_file_{i}", use_container_width=True,help="Redact the file"):
                                    pass
                    
                
                

    # Sample data
    username = "johndoe"
    provider_description = "This is the provider description."

    # Social stats
    seeders = 120
    appearances_in_searches = 300
    completed_enquiries = 50
    reviews = 45
    stars = 4.5
    written_reviews = 30
    mean_engagement = 3.8
    total_interactions = 500
    active_interactions = 150

    # Running agents data
    agents = [
        {
            "name": "Agent 1",
            "skills": ["Skill A", "Skill B"],
            "expertise": "Expertise 1",
            "rating": 4.2,
            "seeders": 60,
            "appearances_in_searches": 150,
            "completed_enquiries": 25,
            "reviews": 20,
            "stars": 4.0,
            "written_reviews": 15,
            "mean_engagement": 3.5,
            "total_interactions": 250,
            "active_interactions": 75
        },
        {
            "name": "Agent 2",
            "skills": ["Skill C", "Skill D"],
            "expertise": "Expertise 2",
            "rating": 4.8,
            "seeders": 60,
            "appearances_in_searches": 150,
            "completed_enquiries": 25,
            "reviews": 25,
            "stars": 4.5,
            "written_reviews": 15,
            "mean_engagement": 4.1,
            "total_interactions": 250,
            "active_interactions": 75
        }
    ]

    with tab3:
        st.text("Welcome!")
        st.subheader(username)
        provider_description = st.text_area("Provider Description", provider_description)
        st.subheader("Social Stats")
        col1, col2, col3 = st.columns(3)
        col1.metric("Seeders", seeders)
        col2.metric("Appearances in Searches", appearances_in_searches)
        col3.metric("Completed Enquiries", completed_enquiries)
        col1.metric("Reviews", reviews)
        col2.metric("Stars", stars)
        col3.metric("Written Reviews", written_reviews)
        col1.metric("Mean Engagement", mean_engagement)
        col2.metric("Total Interactions", total_interactions)
        col3.metric("Active Interactions", active_interactions)
        if st.button("Update Profile", key="update_profile",use_container_width=True):
            pass
        st.subheader("Currently Running Agents")
        with st.container(border=True, height=500):
            for agent in agents:
                with st.expander(agent["name"], expanded=False):
                    st.write(f"**{agent['name']}**")
                    st.write(f"Skills: {', '.join(agent['skills'])}")
                    st.write(f"Expertise: {agent['expertise']}")
                    st.write(f"Rating: {agent['rating']}")
                    st.subheader("Social Stats")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Seeders", agent['seeders'],"-10%")
                    col2.metric("Appearances in Searches", agent['appearances_in_searches'], "-20%")
                    col3.metric("Completed Enquiries", agent['completed_enquiries'], "30%")
                    col1.metric("Reviews", agent['reviews'])
                    col2.metric("Stars", agent['stars'], "10%")
                    col3.metric("Written Reviews", agent['written_reviews'])
                    col1.metric("Mean Engagement", agent['mean_engagement'])
                    col2.metric("Total Interactions", agent['total_interactions'], "-40%")
                    col3.metric("Active Interactions", agent['active_interactions'], "20%")

    # col1, col2, col3 = st.columns(3)
    # col1.metric("Temperature", "70 °F", "1.2 °F")
    # col2.metric("Wind", "9 mph", "-8%")
    # col3.metric("Humidity", "86%", "4%")
    
else:
    "Please log in to continue."
